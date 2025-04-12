import json
import sqlite3
from flask import Flask, render_template, request, jsonify, send_file
from error_classifier import ErrorClassifier
import io
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime

app = Flask(__name__)
classifier = ErrorClassifier()

# Инициализация базы данных SQLite
def init_db():
    conn = sqlite3.connect('cases.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS cases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            personal_data TEXT NOT NULL,
            errors TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# Вызываем инициализацию базы данных при запуске приложения
init_db()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/history')
def history():
    conn = sqlite3.connect('cases.db')
    c = conn.cursor()
    c.execute('SELECT id, personal_data, errors FROM cases')
    rows = c.fetchall()
    conn.close()

    # Преобразуем строки JSON в словари Python
    history = []
    for row in rows:
        case_id, personal_data_json, errors_json = row
        personal_data = json.loads(personal_data_json)
        errors = json.loads(errors_json)
        history.append({
            'id': case_id,
            'personal_data': personal_data,
            'errors': errors
        })

    return render_template('history.html', history=history)

@app.route('/process', methods=['POST'])
def process():
    data = request.form.to_dict()
    
    # Преобразование данных в нужный формат
    case_data = {
        "personal_data": {
            "full_name": data.get('full_name', ''),
            "birth_date": data.get('birth_date', ''),
            "snils": data.get('snils', ''),
            "gender": data.get('gender', ''),
            "citizenship": data.get('citizenship', ''),
            "name_change_info": {
                "old_full_name": data.get('old_full_name', ''),
                "date_changed": data.get('date_changed', '')
            } if data.get('old_full_name') else {},
            "dependents": int(data.get('dependents', 0))
        },
        "work_experience": {
            "total_years": float(data.get('total_years', 0)),
            "records": []
        },
        "pension_points": float(data.get('pension_points', 0)),
        "benefits": data.get('benefits', '').split(',') if data.get('benefits') else [],
        "documents": data.get('documents', '').split(',') if data.get('documents') else [],
        "has_incorrect_document": data.get('has_incorrect_document', 'false').lower() == 'true'
    }

    # Добавление записей о трудовом стаже
    for i in range(1, 6):
        if f'organization_{i}' in data:
            record = {
                "organization": data.get(f'organization_{i}', ''),
                "start_date": data.get(f'start_date_{i}', ''),
                "end_date": data.get(f'end_date_{i}', ''),
                "position": data.get(f'position_{i}', ''),
                "special_conditions": data.get(f'special_conditions_{i}', 'false').lower() == 'true'
            }
            case_data["work_experience"]["records"].append(record)

    # Классификация ошибок
    errors = classifier.classify_errors(case_data)

    # Сохранение дела в базу данных
    conn = sqlite3.connect('cases.db')
    c = conn.cursor()
    c.execute('INSERT INTO cases (personal_data, errors) VALUES (?, ?)',
              (json.dumps(case_data["personal_data"]), json.dumps(errors)))
    conn.commit()
    conn.close()

    return render_template('result.html', personal_data=case_data["personal_data"], errors=errors)

@app.route('/download_document', methods=['POST'])
def download_document():
    # Получение данных из формы
    personal_data_json = request.form.get('personal_data')
    errors_json = request.form.get('errors')
    doc_format = request.form.get('format', 'pdf')  # По умолчанию PDF

    personal_data = json.loads(personal_data_json)
    errors = json.loads(errors_json)

    # Маскировка персональных данных
    masked_data = {
        "full_name": "ФИО скрыто",
        "birth_date": "**.**.****",
        "snils": "***-***-*** **",
        "gender": personal_data["gender"],
        "citizenship": "Гражданство скрыто",
        "name_change_info": {
            "old_full_name": "ФИО скрыто",
            "date_changed": personal_data["name_change_info"]["date_changed"]
        } if personal_data["name_change_info"] else {},
        "dependents": "Данные скрыты"
    }

    # Текущая дата
    current_date = datetime.now().strftime("%d.%m.%Y")

    if doc_format == 'pdf':
        # Создание PDF-документа
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=30, bottomMargin=30)
        styles = getSampleStyleSheet()
        elements = []

        # Заголовок
        title_style = ParagraphStyle(
            name='TitleStyle',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=20,
            alignment=1  # Центрирование
        )
        elements.append(Paragraph("Решение по пенсионному делу", title_style))

        # Дата
        date_style = ParagraphStyle(
            name='DateStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=10,
            alignment=2  # Выравнивание вправо
        )
        elements.append(Paragraph(f"Дата: {current_date}", date_style))

        # Замаскированные персональные данные
        elements.append(Paragraph("Персональные данные", styles['Heading2']))
        normal_style = ParagraphStyle(
            name='NormalStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=6
        )
        elements.append(Paragraph(f"ФИО: {masked_data['full_name']}", normal_style))
        elements.append(Paragraph(f"Дата рождения: {masked_data['birth_date']}", normal_style))
        elements.append(Paragraph(f"СНИЛС: {masked_data['snils']}", normal_style))
        elements.append(Paragraph(f"Пол: {masked_data['gender']}", normal_style))
        elements.append(Paragraph(f"Гражданство: {masked_data['citizenship']}", normal_style))
        if masked_data["name_change_info"]:
            elements.append(Paragraph(f"Смена имени: {masked_data['name_change_info']['old_full_name']} (дата: {masked_data['name_change_info']['date_changed']})", normal_style))
        elements.append(Paragraph(f"Иждивенцы: {masked_data['dependents']}", normal_style))

        elements.append(Spacer(1, 12))

        # Обоснование отказа
        elements.append(Paragraph("Решение по делу", styles['Heading2']))
        if errors:
            elements.append(Paragraph("На основании проведённого анализа в предоставлении пенсии отказано по следующим причинам:", normal_style))
            elements.append(Spacer(1, 12))
            elements.append(Paragraph("Выявленные ошибки:", styles['Heading3']))
            
            for error in errors:
                error_text = (
                    f"{error['code']}: {error['description']}<br/>"
                    f"Закон: {error['law']}<br/>"
                    f"Рекомендация: {error['recommendation']}"
                )
                elements.append(Paragraph(error_text, normal_style))
                elements.append(Spacer(1, 6))
        else:
            elements.append(Paragraph("Ошибок не выявлено. Пенсия может быть предоставлена.", normal_style))

        # Генерация PDF
        doc.build(elements)
        buffer.seek(0)

        # Отправка файла пользователю
        return send_file(
            buffer,
            as_attachment=True,
            download_name="pension_decision.pdf",
            mimetype="application/pdf"
        )

    elif doc_format == 'docx':
        # Создание DOCX-документа
        doc = Document()
        
        # Заголовок
        title = doc.add_heading("Решение по пенсионному делу", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата
        date_p = doc.add_paragraph(f"Дата: {current_date}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.runs[0].font.size = Pt(12)

        # Замаскированные персональные данные
        doc.add_heading("Персональные данные", level=2)
        p = doc.add_paragraph(f"ФИО: {masked_data['full_name']}")
        p.runs[0].font.size = Pt(12)
        p = doc.add_paragraph(f"Дата рождения: {masked_data['birth_date']}")
        p.runs[0].font.size = Pt(12)
        p = doc.add_paragraph(f"СНИЛС: {masked_data['snils']}")
        p.runs[0].font.size = Pt(12)
        p = doc.add_paragraph(f"Пол: {masked_data['gender']}")
        p.runs[0].font.size = Pt(12)
        p = doc.add_paragraph(f"Гражданство: {masked_data['citizenship']}")
        p.runs[0].font.size = Pt(12)
        if masked_data["name_change_info"]:
            p = doc.add_paragraph(f"Смена имени: {masked_data['name_change_info']['old_full_name']} (дата: {masked_data['name_change_info']['date_changed']})")
            p.runs[0].font.size = Pt(12)
        p = doc.add_paragraph(f"Иждивенцы: {masked_data['dependents']}")
        p.runs[0].font.size = Pt(12)

        # Обоснование отказа
        doc.add_heading("Решение по делу", level=2)
        if errors:
            p = doc.add_paragraph("На основании проведённого анализа в предоставлении пенсии отказано по следующим причинам:")
            p.runs[0].font.size = Pt(12)
            doc.add_heading("Выявленные ошибки:", level=3)
            for error in errors:
                p = doc.add_paragraph(f"{error['code']}: {error['description']}")
                p.runs[0].font.size = Pt(12)
                p = doc.add_paragraph(f"Закон: {error['law']}")
                p.runs[0].font.size = Pt(12)
                p = doc.add_paragraph(f"Рекомендация: {error['recommendation']}")
                p.runs[0].font.size = Pt(12)
                doc.add_paragraph()  # Пустая строка для разделения
        else:
            p = doc.add_paragraph("Ошибок не выявлено. Пенсия может быть предоставлена.")
            p.runs[0].font.size = Pt(12)

        # Сохранение DOCX в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Отправка файла пользователю
        return send_file(
            buffer,
            as_attachment=True,
            download_name="pension_decision.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        return "Неподдерживаемый формат документа", 400

if __name__ == "__main__":
    app.run(debug=True)