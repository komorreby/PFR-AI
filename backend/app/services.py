import io
import asyncio
from datetime import datetime, date
from typing import List, Dict, Any, Tuple, Optional

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import cm

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn # Для установки кастомного шрифта для всего документа

# Попытка импортировать markdown2. Если не получится, будет использоваться заглушка.
MarkdownImportError = None
try:
    import markdown2
except ImportError as e:
    MarkdownImportError = e
    markdown2 = None
    # Добавим логирование, если markdown2 не найден
    import logging
    logging.warning("Markdown2 library not found. Markdown formatting will be basic. Please install with: pip install markdown2")

from .models import DocumentFormat # Используем Enum и модель ошибки
# Добавляем импорты для конфигурационных моделей, если они понадобятся для типизации
from .config_models.config_models import PensionTypeInfo, PensionTypeDocuments, DocumentDetail

def mask_personal_data(personal_data: Dict[str, Any]) -> Dict[str, Any]:
    """Маскирует чувствительные персональные данные."""
    masked_data = {
        "full_name": "[ФИО СКРЫТО]",
        "birth_date": "[ДАТА РОЖДЕНИЯ СКРЫТА]", # Используем более общую маску
        "snils": "[СНИЛС СКРЫТ]",
        "gender": personal_data.get("gender", "[ПОЛ СКРЫТ]"), # Если пол важен, он должен быть в personal_data
        "citizenship": "[ГРАЖДАНСТВО СКРЫТО]",
        "name_change_info": {
            "old_full_name": "[ПРЕЖНЕЕ ФИО СКРЫТО]",
            "date_changed": "[ДАТА СМЕНЫ ФИО СКРЫТА]"
        } if personal_data and personal_data.get("name_change_info") else {}, # Проверка на None для personal_data
        "dependents": "[КОЛИЧЕСТВО ИЖДИВЕНЦЕВ СКРЫТО]" # Маскируем количество
    }
    return masked_data

def _get_pension_type_display_name(pension_type_id: str, pension_types_config: List[PensionTypeInfo]) -> str:
    if pension_types_config:
        for pt in pension_types_config:
            if pt.id == pension_type_id:
                return pt.display_name
    return pension_type_id # Возвращаем ID, если имя не найдено

def _get_document_display_name(doc_id: str, doc_requirements_config: Dict[str, PensionTypeDocuments], pension_type_id_for_context: Optional[str] = None) -> str:
    # Сначала ищем в контексте конкретного типа пенсии, если он предоставлен
    if pension_type_id_for_context and doc_requirements_config and pension_type_id_for_context in doc_requirements_config:
        for doc_detail in doc_requirements_config[pension_type_id_for_context].documents:
            if doc_detail.id == doc_id:
                return doc_detail.name
    # Если не найдено или контекст не дан, ищем по всем типам
    if doc_requirements_config:
        for pt_id, reqs in doc_requirements_config.items():
            for doc_detail in reqs.documents:
                if doc_detail.id == doc_id:
                    return doc_detail.name
    return doc_id # Возвращаем ID, если имя не найдено

def _convert_markdown_to_html_for_reportlab(md_text: str) -> str:
    """Конвертирует Markdown в HTML, подходящий для ReportLab Paragraph."""
    # Добавим логирование входного и выходного текста
    import logging
    logger_services = logging.getLogger(__name__) # Используем логгер текущего модуля
    logger_services.debug(f"Original Markdown for PDF: {md_text!r}")

    if markdown2:
        html = markdown2.markdown(md_text, extras=["break-on-newline", "cuddled-lists", "smarty-pants"])
        logger_services.debug(f"Converted HTML (using markdown2) for PDF: {html!r}")
        return html
    else:
        logger_services.warning("markdown2 not available, using basic HTML escaping for PDF.")
        import html as html_converter
        escaped_text = html_converter.escape(md_text)
        html_fallback = escaped_text.replace("\n", "<br/>")
        logger_services.debug(f"Fallback HTML for PDF: {html_fallback!r}")
        return html_fallback

def _strip_markdown_for_docx(md_text: str) -> str:
    """Удаляет Markdown разметку из текста (упрощенный вариант)."""
    if markdown2:
        html_text = markdown2.markdown(md_text)
        import re
        clean_text = re.sub(r'<[^>]+>', '', html_text) # Удаляем HTML теги
        clean_text = clean_text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        return clean_text
    else:
        text = md_text.replace("**", "").replace("__", "")\
                      .replace("*", "").replace("_", "")\
                      .replace("\n", " ") # Заменяем переносы на пробел для DOCX
        # Это очень грубо и не учтет списки, заголовки и т.д.
        return text

def _generate_pdf_report(
    case_details: Dict[str, Any],
    pension_types_list_config: List[PensionTypeInfo],
    doc_requirements_config: Dict[str, PensionTypeDocuments]
) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=2*cm, rightMargin=1.5*cm,
                            topMargin=1.5*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elements = []

    # Шрифты
    # Попробуем стандартный шрифт Helvetica, который обычно хорошо поддерживается
    # и часто включает кириллицу. Если нет, понадобится явная регистрация TTF.
    font_name = 'Helvetica'
    bold_font_name = 'Helvetica-Bold' # Стандартное имя для жирного Helvetica
    
    # Применяем базовый шрифт к стандартным стилям
    styles['Normal'].fontName = font_name
    styles['Heading1'].fontName = bold_font_name # Заголовки делаем жирными
    styles['Heading2'].fontName = bold_font_name
    styles['Heading3'].fontName = bold_font_name
    styles['Bullet'].fontName = font_name
    styles['Definition'].fontName = font_name

    # Основные Стили
    title_style = ParagraphStyle(name='TitleStyle', parent=styles['Heading1'], fontSize=16, alignment=TA_CENTER, spaceAfter=12, fontName=bold_font_name)
    doc_date_style = ParagraphStyle(name='DocDateStyle', parent=styles['Normal'], fontSize=10, alignment=TA_RIGHT, spaceAfter=10, fontName=font_name)
    system_name_style = ParagraphStyle(name='SystemNameStyle', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, spaceBefore=6, spaceAfter=12, fontName=font_name)
    
    section_title_style = ParagraphStyle(name='SectionTitle', parent=styles['Heading2'], fontSize=14, spaceBefore=12, spaceAfter=6, fontName=bold_font_name)
    subsection_title_style = ParagraphStyle(name='SubSectionTitle', parent=styles['Heading3'], fontSize=12, spaceBefore=8, spaceAfter=4, fontName=bold_font_name)
    
    normal_style_justify = ParagraphStyle(name='NormalJustify', parent=styles['Normal'], fontSize=12, alignment=TA_JUSTIFY, spaceAfter=6, fontName=font_name)
    normal_style_left = ParagraphStyle(name='NormalLeft', parent=styles['Normal'], fontSize=12, alignment=TA_LEFT, spaceAfter=6, fontName=font_name)
    # Для label_style используем <b /> теги внутри текста параграфа, ReportLab должен их обработать с базовым шрифтом
    label_style = ParagraphStyle(name='LabelStyle', parent=normal_style_left, fontSize=12, fontName=font_name) 
    
    error_style = ParagraphStyle(name='ErrorStyle', parent=normal_style_justify, textColor=colors.red, spaceAfter=6, fontName=font_name)
    footer_style = ParagraphStyle(name='FooterStyle', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, spaceBefore=24, fontName=font_name)

    # Извлечение данных из case_details
    case_id = case_details.get("id", "N/A")
    current_date_str = datetime.now().strftime("%d.%m.%Y")
    
    personal_data = case_details.get("personal_data", {})
    masked_personal_data = mask_personal_data(personal_data if personal_data else {})

    # Шапка документа
    elements.append(Paragraph(f"РЕШЕНИЕ № {case_id}", title_style))
    elements.append(Paragraph(f"по заявлению о назначении пенсии", styles['Normal'])) # Подзаголовок если нужен
    elements.append(Paragraph(f"Дата формирования: {current_date_str}", doc_date_style))
    elements.append(Spacer(1, 0.5*cm))

    # Раздел 1: Информация о заявителе (обезличенная)
    elements.append(Paragraph("1. Сведения о заявителе (обезличенные)", section_title_style))
    elements.append(Paragraph(f"<b>Фамилия, имя, отчество (при наличии):</b> {masked_personal_data['full_name']}", label_style))
    elements.append(Paragraph(f"<b>Дата рождения:</b> {masked_personal_data['birth_date']}", label_style))
    elements.append(Paragraph(f"<b>СНИЛС:</b> {masked_personal_data['snils']}", label_style))
    elements.append(Paragraph(f"<b>Пол:</b> {masked_personal_data['gender']}", label_style))
    elements.append(Paragraph(f"<b>Гражданство:</b> {masked_personal_data['citizenship']}", label_style))
    if masked_personal_data.get("name_change_info") and masked_personal_data["name_change_info"].get('old_full_name') != "[ПРЕЖНЕЕ ФИО СКРЫТО]":
        elements.append(Paragraph("<b>Сведения о ранее измененном ФИО:</b>", label_style))
        elements.append(Paragraph(f"  Прежнее ФИО: {masked_personal_data['name_change_info']['old_full_name']}", label_style))
        elements.append(Paragraph(f"  Дата изменения: {masked_personal_data['name_change_info']['date_changed']}", label_style))
    elements.append(Paragraph(f"<b>Количество заявленных иждивенцев:</b> {masked_personal_data['dependents']}", label_style))
    elements.append(Spacer(1, 0.5*cm))

    # Раздел 2: Запрашиваемый вид пенсионного обеспечения
    elements.append(Paragraph("2. Запрашиваемый вид пенсионного обеспечения", section_title_style))
    pension_type_id = case_details.get("pension_type", "Не указан")
    pension_type_name = _get_pension_type_display_name(pension_type_id, pension_types_list_config)
    elements.append(Paragraph(pension_type_name, normal_style_justify))
    elements.append(Spacer(1, 0.5*cm))

    # Раздел 3: Представленные сведения и документы
    elements.append(Paragraph("3. Представленные сведения и документы", section_title_style))
    
    disability_info = case_details.get("disability")
    if disability_info:
        elements.append(Paragraph("3.1. Сведения об инвалидности", subsection_title_style))
        elements.append(Paragraph(f"<b>Группа инвалидности:</b> {disability_info.get('group', 'Не указана')}", label_style))
        dis_date = disability_info.get('date')
        dis_date_str = dis_date.strftime("%d.%m.%Y") if isinstance(dis_date, datetime) or isinstance(dis_date, date) else str(dis_date or 'Не указана')
        elements.append(Paragraph(f"<b>Дата установления:</b> {dis_date_str}", label_style))
        elements.append(Paragraph(f"<b>Номер справки МСЭ:</b> {disability_info.get('cert_number', 'Не указан')}", label_style))

    work_experience_info = case_details.get("work_experience")
    pension_points = case_details.get("pension_points")
    if work_experience_info or pension_points is not None:
        elements.append(Paragraph("3.2. Сведения о трудовом стаже и пенсионных баллах", subsection_title_style))
        if work_experience_info:
            elements.append(Paragraph(f"<b>Общий страховой стаж (лет):</b> {work_experience_info.get('total_years', 'Не указан')}", label_style))
        if pension_points is not None:
            elements.append(Paragraph(f"<b>Индивидуальный пенсионный коэффициент (ИПК):</b> {pension_points}", label_style))
        
        if work_experience_info and work_experience_info.get("records"):
            elements.append(Paragraph("<b>Периоды трудовой деятельности:</b>", label_style))
            table_data = [["Организация", "Должность", "Период работы", "Особые условия"]]
            for record in work_experience_info["records"]:
                start_date_str = record.get('start_date').strftime("%d.%m.%Y") if record.get('start_date') else 'N/A'
                end_date_str = record.get('end_date').strftime("%d.%m.%Y") if record.get('end_date') else 'N/A'
                period_str = f"{start_date_str} - {end_date_str}"
                table_data.append([
                    Paragraph(record.get('organization', 'N/A')  + " (Данные могут быть обезличены)", styles['Normal']), # Обезличивание если нужно
                    Paragraph(record.get('position', 'N/A'), styles['Normal']),
                    Paragraph(period_str, styles['Normal']),
                    Paragraph("Да" if record.get('special_conditions') else "Нет", styles['Normal'])
                ])
            
            if len(table_data) > 1: # Если есть записи кроме заголовка
                work_table = Table(table_data, colWidths=[6*cm, 4*cm, 4*cm, 3*cm])
                work_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
                    ('ALIGN',(0,0),(-1,-1),'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), bold_font_name), # Используем bold_font_name
                    ('FONTNAME', (0,1), (-1,-1), font_name), # Обычный шрифт для данных
                    ('FONTSIZE', (0,0), (-1,-1), 10),
                    ('BOTTOMPADDING', (0,0), (-1,0), 10),
                    ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ]))
                elements.append(work_table)
            else:
                elements.append(Paragraph("Записи о периодах трудовой деятельности отсутствуют.", normal_style_left))

    benefits = case_details.get("benefits")
    if benefits:
        elements.append(Paragraph("3.3. Заявленные льготы", subsection_title_style))
        for benefit in benefits:
            elements.append(Paragraph(f"• {benefit}", styles['Bullet']))
    
    submitted_documents = case_details.get("submitted_documents")
    if submitted_documents:
        elements.append(Paragraph("3.4. Перечень представленных заявителем документов", subsection_title_style))
        for doc_id in submitted_documents:
            doc_name = _get_document_display_name(doc_id, doc_requirements_config, pension_type_id)
            elements.append(Paragraph(f"• {doc_name} (ID: {doc_id})", styles['Bullet']))

    if case_details.get("has_incorrect_document"):
        elements.append(Paragraph("3.5. Отметка о некорректно оформленных документах", subsection_title_style))
        elements.append(Paragraph("Заявителем отмечено наличие некорректно оформленных документов.", normal_style_left))

    other_documents_data = case_details.get("other_documents_extracted_data")
    if other_documents_data:
        elements.append(Paragraph("3.6. Сведения из дополнительно загруженных документов (по результатам OCR)", subsection_title_style))
        for i, ocr_doc in enumerate(other_documents_data):
            elements.append(Paragraph(f"<b>Документ {i+1}:</b>", label_style))
            doc_type_display = ocr_doc.get("standardized_document_type") or ocr_doc.get("identified_document_type") or "Тип не определен"
            elements.append(Paragraph(f"  <i>Тип документа (определенный системой):</i> {doc_type_display}", normal_style_left))
            extracted_fields = ocr_doc.get("extracted_fields")
            if extracted_fields and isinstance(extracted_fields, dict):
                elements.append(Paragraph("  <i>Ключевые извлеченные поля (обезличенные):</i>", normal_style_left))
                for key, val in extracted_fields.items():
                     # Простое обезличивание для примера, может потребоваться более сложная логика
                    val_display = "[СКРЫТО]" if isinstance(val, str) and len(val) > 3 else str(val)
                    elements.append(Paragraph(f"    - {key}: {val_display}", normal_style_left))
            multimodal_assessment = ocr_doc.get("multimodal_assessment")
            if multimodal_assessment:
                 elements.append(Paragraph(f"  <i>Оценка документа системой:</i> {multimodal_assessment}", normal_style_left))
        elements.append(Spacer(1, 0.2*cm))

    elements.append(Spacer(1, 0.5*cm))

    # Раздел 4: Результаты автоматизированного анализа
    elements.append(Paragraph("4. Результаты автоматизированного анализа и решение", section_title_style))
    final_status = case_details.get("final_status", "Статус не определен")
    status_display = final_status
    if final_status == "СООТВЕТСТВУЕТ":
        status_display = "Право на назначение пенсии подтверждено"
    elif final_status == "НЕ СООТВЕТСТВУЕТ":
        status_display = "В праве на назначение пенсии отказано (условия для назначения пенсии не выполнены)"
    elif final_status == "PROCESSING":
        status_display = "Дело находится в обработке"
    elif final_status == "ERROR_PROCESSING":
        status_display = "Ошибка при обработке дела"
        
    elements.append(Paragraph("4.1. Итоговое решение системы", subsection_title_style))
    elements.append(Paragraph(f"<b>{status_display}</b>", normal_style_justify))

    elements.append(Paragraph("4.2. Обоснование решения", subsection_title_style))
    explanation_md = case_details.get("final_explanation", "Обоснование отсутствует.")
    explanation_html_for_pdf = _convert_markdown_to_html_for_reportlab(explanation_md)
    elements.append(Paragraph(explanation_html_for_pdf, normal_style_justify))

    rag_confidence = case_details.get("rag_confidence")
    if rag_confidence is not None:
        elements.append(Paragraph("4.3. Степень уверенности системы в принятом решении", subsection_title_style))
        elements.append(Paragraph(f"{rag_confidence*100:.1f}%", normal_style_justify))
    elements.append(Spacer(1, 0.5*cm))
    
    # Раздел 5: Выявленные ошибки/несоответствия (если есть)
    errors_list = case_details.get("errors", [])
    if errors_list:
        elements.append(Paragraph("5. Выявленные ошибки/несоответствия", section_title_style))
        for err_idx, error_item in enumerate(errors_list):
            elements.append(Paragraph(f"<b>Ошибка {err_idx + 1}:</b>", subsection_title_style))
            elements.append(Paragraph(f"  Код: {error_item.get('code', 'N/A')}", normal_style_left))
            elements.append(Paragraph(f"  Описание: {error_item.get('description', 'N/A')}", normal_style_left))
            if error_item.get('law'):
                elements.append(Paragraph(f"  Основание (закон): {error_item.get('law')}", normal_style_left))
            if error_item.get('recommendation'):
                elements.append(Paragraph(f"  Рекомендация: {error_item.get('recommendation')}", normal_style_left))
        elements.append(Spacer(1, 0.5*cm))

    # Подвал документа
    elements.append(Paragraph("Сформировано автоматизированной системой анализа пенсионных дел 'PFR-AI'.", footer_style))
    elements.append(Paragraph("Данное решение носит предварительный характер.", ParagraphStyle(name='FooterDisclaimer', parent=footer_style, fontSize=9)))


    # Сборка документа и пагинация
    def add_page_numbers(canvas, doc_template):
        canvas.saveState()
        canvas.setFont(font_name, 9) # Используем базовый font_name для номеров страниц
        page_num_text = f"Страница {doc_template.page}"
        canvas.drawCentredString(A4[0]/2, 1*cm, page_num_text)
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_page_numbers, onLaterPages=add_page_numbers)
    buffer.seek(0)
    return buffer

def _generate_docx_report(
    case_details: Dict[str, Any],
    pension_types_list_config: List[PensionTypeInfo],
    doc_requirements_config: Dict[str, PensionTypeDocuments]
) -> io.BytesIO:
    doc = Document()
    # Установка шрифта по умолчанию для всего документа (Times New Roman)
    # Это более надежный способ, чем менять стиль 'Normal' для каждого параграфа
    doc_element = doc.element.body
    if doc_element is not None:
        # Получаем или создаем секцию свойств документа
        sectPr = doc_element.find(qn('w:sectPr'))
        if sectPr is None:
            sectPr = doc_element.makeelement(qn('w:sectPr'))
            doc_element.append(sectPr)
        # Устанавливаем основной шрифт для всего документа (если возможно через стили)
        # Для DOCX лучше всего определить стили в шаблоне, если это возможно.
        # Программно можно менять стиль 'Normal' или применять шрифт к каждому Run.
        # Попробуем установить стиль 'Normal'
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Извлечение данных
    case_id = case_details.get("id", "N/A")
    current_date_str = datetime.now().strftime("%d.%m.%Y")
    personal_data = case_details.get("personal_data", {})
    masked_personal_data = mask_personal_data(personal_data if personal_data else {})

    # Шапка документа
    title = doc.add_heading(f"РЕШЕНИЕ № {case_id}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"по заявлению о назначении пенсии", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_p = doc.add_paragraph(f"Дата формирования: {current_date_str}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph() # Spacer

    # Раздел 1: Информация о заявителе (обезличенная)
    doc.add_heading("1. Сведения о заявителе (обезличенные)", level=2)
    p = doc.add_paragraph()
    p.add_run("Фамилия, имя, отчество (при наличии): ").bold = True
    p.add_run(masked_personal_data['full_name'])
    p = doc.add_paragraph()
    p.add_run("Дата рождения: ").bold = True
    p.add_run(masked_personal_data['birth_date'])
    p = doc.add_paragraph()
    p.add_run("СНИЛС: ").bold = True
    p.add_run(masked_personal_data['snils'])
    p = doc.add_paragraph()
    p.add_run("Пол: ").bold = True
    p.add_run(masked_personal_data['gender'])
    p = doc.add_paragraph()
    p.add_run("Гражданство: ").bold = True
    p.add_run(masked_personal_data['citizenship'])
    
    if masked_personal_data.get("name_change_info") and masked_personal_data["name_change_info"].get('old_full_name') != "[ПРЕЖНЕЕ ФИО СКРЫТО]":
        p = doc.add_paragraph()
        p.add_run("Сведения о ранее измененном ФИО:").bold = True
        sub_p = doc.add_paragraph(f"  Прежнее ФИО: {masked_personal_data['name_change_info']['old_full_name']}", style='ListBullet') # или просто отступ
        sub_p.paragraph_format.left_indent = Cm(0.5)
        sub_p = doc.add_paragraph(f"  Дата изменения: {masked_personal_data['name_change_info']['date_changed']}", style='ListBullet')
        sub_p.paragraph_format.left_indent = Cm(0.5)
        
    p = doc.add_paragraph()
    p.add_run("Количество заявленных иждивенцев: ").bold = True
    p.add_run(masked_personal_data['dependents'])
    doc.add_paragraph()

    # Раздел 2: Запрашиваемый вид пенсионного обеспечения
    doc.add_heading("2. Запрашиваемый вид пенсионного обеспечения", level=2)
    pension_type_id = case_details.get("pension_type", "Не указан")
    pension_type_name = _get_pension_type_display_name(pension_type_id, pension_types_list_config)
    doc.add_paragraph(pension_type_name, style='Normal')
    doc.add_paragraph()

    # Раздел 3: Представленные сведения и документы
    doc.add_heading("3. Представленные сведения и документы", level=2)
    
    disability_info = case_details.get("disability")
    if disability_info:
        doc.add_heading("3.1. Сведения об инвалидности", level=3)
        p = doc.add_paragraph()
        p.add_run("Группа инвалидности: ").bold = True
        p.add_run(str(disability_info.get('group', 'Не указана')))
        dis_date = disability_info.get('date')
        dis_date_str = dis_date.strftime("%d.%m.%Y") if isinstance(dis_date, datetime) or isinstance(dis_date, date) else str(dis_date or 'Не указана')
        p = doc.add_paragraph()
        p.add_run("Дата установления: ").bold = True
        p.add_run(dis_date_str)
        p = doc.add_paragraph()
        p.add_run("Номер справки МСЭ: ").bold = True
        p.add_run(str(disability_info.get('cert_number', 'Не указан')))

    work_experience_info = case_details.get("work_experience")
    pension_points = case_details.get("pension_points")
    if work_experience_info or pension_points is not None:
        doc.add_heading("3.2. Сведения о трудовом стаже и пенсионных баллах", level=3)
        if work_experience_info:
            p = doc.add_paragraph()
            p.add_run("Общий страховой стаж (лет): ").bold = True
            p.add_run(str(work_experience_info.get('total_years', 'Не указан')))
        if pension_points is not None:
            p = doc.add_paragraph()
            p.add_run("Индивидуальный пенсионный коэффициент (ИПК): ").bold = True
            p.add_run(str(pension_points))
        
        if work_experience_info and work_experience_info.get("records"):
            p = doc.add_paragraph()
            p.add_run("Периоды трудовой деятельности:").bold = True
            
            if work_experience_info["records"]:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid' # Применяем стиль таблицы
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Организация"
                hdr_cells[1].text = "Должность"
                hdr_cells[2].text = "Период работы"
                hdr_cells[3].text = "Особые условия"
                for cell in hdr_cells: # Делаем заголовки жирными
                    for paragraph in cell.paragraphs:
                        for run_ in paragraph.runs:
                            run_.bold = True
                
                for record in work_experience_info["records"]:
                    row_cells = table.add_row().cells
                    start_date_str = record.get('start_date').strftime("%d.%m.%Y") if record.get('start_date') else 'N/A'
                    end_date_str = record.get('end_date').strftime("%d.%m.%Y") if record.get('end_date') else 'N/A'
                    period_str = f"{start_date_str} - {end_date_str}"
                    row_cells[0].text = record.get('organization', 'N/A') + " (Данные могут быть обезличены)"
                    row_cells[1].text = record.get('position', 'N/A')
                    row_cells[2].text = period_str
                    row_cells[3].text = "Да" if record.get('special_conditions') else "Нет"
            else:
                doc.add_paragraph("Записи о периодах трудовой деятельности отсутствуют.", style='Normal')


    benefits = case_details.get("benefits")
    if benefits:
        doc.add_heading("3.3. Заявленные льготы", level=3)
        for benefit in benefits:
            doc.add_paragraph(str(benefit), style='ListBullet')
    
    submitted_documents = case_details.get("submitted_documents")
    if submitted_documents:
        doc.add_heading("3.4. Перечень представленных заявителем документов", level=3)
        for doc_id in submitted_documents:
            doc_name = _get_document_display_name(doc_id, doc_requirements_config, pension_type_id)
            doc.add_paragraph(f"{doc_name} (ID: {doc_id})", style='ListBullet')

    if case_details.get("has_incorrect_document"):
        doc.add_heading("3.5. Отметка о некорректно оформленных документах", level=3)
        doc.add_paragraph("Заявителем отмечено наличие некорректно оформленных документов.", style='Normal')

    other_documents_data = case_details.get("other_documents_extracted_data")
    if other_documents_data:
        doc.add_heading("3.6. Сведения из дополнительно загруженных документов (по результатам OCR)", level=3)
        for i, ocr_doc in enumerate(other_documents_data):
            p = doc.add_paragraph()
            p.add_run(f"Документ {i+1}:").bold = True
            doc_type_display = ocr_doc.get("standardized_document_type") or ocr_doc.get("identified_document_type") or "Тип не определен"
            sub_p = doc.add_paragraph(f"  Тип документа (определенный системой): {doc_type_display}", style='Normal')
            sub_p.paragraph_format.left_indent = Cm(0.5)
            extracted_fields = ocr_doc.get("extracted_fields")
            if extracted_fields and isinstance(extracted_fields, dict):
                sub_p = doc.add_paragraph("  Ключевые извлеченные поля (обезличенные):", style='Normal')
                sub_p.paragraph_format.left_indent = Cm(0.5)
                for key, val in extracted_fields.items():
                    val_display = "[СКРЫТО]" if isinstance(val, str) and len(val) > 3 else str(val)
                    field_p = doc.add_paragraph(f"    - {key}: {val_display}", style='ListBullet') # Или Normal с отступом
                    field_p.paragraph_format.left_indent = Cm(1.0)
            multimodal_assessment = ocr_doc.get("multimodal_assessment")
            if multimodal_assessment:
                 sub_p = doc.add_paragraph(f"  Оценка документа системой: {multimodal_assessment}", style='Normal')
                 sub_p.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

    # Раздел 4: Результаты автоматизированного анализа
    doc.add_heading("4. Результаты автоматизированного анализа и решение", level=2)
    final_status = case_details.get("final_status", "Статус не определен")
    status_display = final_status
    if final_status == "СООТВЕТСТВУЕТ":
        status_display = "Право на назначение пенсии подтверждено"
    elif final_status == "НЕ СООТВЕТСТВУЕТ":
        status_display = "В праве на назначение пенсии отказано (условия для назначения пенсии не выполнены)"
    elif final_status == "PROCESSING":
        status_display = "Дело находится в обработке"
    elif final_status == "ERROR_PROCESSING":
        status_display = "Ошибка при обработке дела"
        
    doc.add_heading("4.1. Итоговое решение системы", level=3)
    p = doc.add_paragraph()
    p.add_run(status_display).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("4.2. Обоснование решения", level=3)
    explanation = case_details.get("final_explanation", "Обоснование отсутствует.")
    doc.add_paragraph(explanation, style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rag_confidence = case_details.get("rag_confidence")
    if rag_confidence is not None:
        doc.add_heading("4.3. Степень уверенности системы в принятом решении", level=3)
        doc.add_paragraph(f"{rag_confidence*100:.1f}%", style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Раздел 5: Выявленные ошибки/несоответствия (если есть)
    errors_list = case_details.get("errors", [])
    if errors_list:
        doc.add_heading("5. Выявленные ошибки/несоответствия", level=2) # Сделаем Level 2 для выделения
        for err_idx, error_item in enumerate(errors_list):
            p_err_title = doc.add_paragraph()
            p_err_title.add_run(f"Ошибка {err_idx + 1}:").bold = True
            
            doc.add_paragraph(f"  Код: {error_item.get('code', 'N/A')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
            doc.add_paragraph(f"  Описание: {error_item.get('description', 'N/A')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
            if error_item.get('law'):
                doc.add_paragraph(f"  Основание (закон): {error_item.get('law')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
            if error_item.get('recommendation'):
                doc.add_paragraph(f"  Рекомендация: {error_item.get('recommendation')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    # Подвал документа
    doc.add_paragraph() # Spacer
    footer_p1 = doc.add_paragraph("Сформировано автоматизированной системой анализа пенсионных дел 'PFR-AI'.")
    footer_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p1.runs: run.font.size = Pt(10)
    
    footer_p2 = doc.add_paragraph("Данное решение носит предварительный характер.")
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p2.runs: run.font.size = Pt(9)


    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

async def generate_document(
    case_details: Dict[str, Any], # Изменен параметр
    pension_types_list_config: List[PensionTypeInfo], # Добавлен параметр
    doc_requirements_config: Dict[str, PensionTypeDocuments], # Добавлен параметр
    document_format: DocumentFormat
) -> Tuple[io.BytesIO, str, str]:
    """Асинхронно генерирует документ указанного формата, используя asyncio.to_thread."""

    # current_date больше не нужен как параметр, генерируется внутри _generate_X_report
    # masked_data также генерируется внутри _generate_X_report из case_details["personal_data"]
    # errors также из case_details
    # pension_type, final_status, explanation, case_id также из case_details

    doc_format_value = document_format.value

    if doc_format_value == DocumentFormat.pdf.value:
        buffer = await asyncio.to_thread(
            _generate_pdf_report, 
            case_details, pension_types_list_config, doc_requirements_config
        )
        filename = f"pension_decision_{case_details.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        mimetype = "application/pdf"
    elif doc_format_value == DocumentFormat.docx.value:
        buffer = await asyncio.to_thread(
            _generate_docx_report, 
            case_details, pension_types_list_config, doc_requirements_config
        )
        filename = f"pension_decision_{case_details.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise ValueError(f"Unsupported document format: {doc_format_value}")

    return buffer, filename, mimetype
